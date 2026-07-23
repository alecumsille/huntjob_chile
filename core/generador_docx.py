"""
Módulo de compilación de CVs en formato Microsoft Word (.docx).
Mantiene el mismo esqueleto ATS estandarizado y renderiza en negrita
todas las palabras clave y requisitos en etiquetas <b>...</b> o **...**.
"""

import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor

from core.generador_pdf import construir_secciones_cv

COLOR_TITULO = RGBColor(27, 54, 93)      # Azul Marino Ejecutivo
COLOR_SUBTITULO = RGBColor(0, 86, 145)   # Azul secundario
COLOR_TEXTO = RGBColor(51, 51, 51)       # Gris oscuro #333333
COLOR_CONTACTO = RGBColor(119, 119, 119)  # Gris suave


def _agregar_parrafo_con_negritas(doc: Document, texto: str, es_bullet: bool = False):
    """Agrega un párrafo a python-docx interpretando tags <b>...</b> o **...** como texto en negrita."""
    if es_bullet:
        p = doc.add_paragraph(style="List Bullet")
    else:
        p = doc.add_paragraph()

    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15

    # Convertir markdown **text** a <b>text</b>
    texto_norm = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", texto.strip())
    # Separar fragmentos entre <b> y </b>
    partes = re.split(r"(<b>.*?</b>)", texto_norm)

    for parte in partes:
        if not parte:
            continue
        if parte.startswith("<b>") and parte.endswith("</b>"):
            contenido = parte[3:-4]
            run = p.add_run(contenido)
            run.bold = True
            run.font.color.rgb = COLOR_TEXTO
        else:
            # Limpiar cualquier tag HTML residual
            limpio = re.sub(r"<[^>]+>", "", parte)
            if limpio:
                run = p.add_run(limpio)
                run.font.color.rgb = COLOR_TEXTO


def generar_docx_cv(
    perfil: dict,
    resumen_profesional: str,
    bullets_por_trabajo: list[list[str]],
    puesto: str,
    resumen_fit: list[dict] | None = None,
) -> bytes:
    """
    Compila un CV en formato Word (.docx) a partir del esqueleto ATS estándar.
    Retorna un buffer de bytes listo para su descarga.
    """
    secciones = construir_secciones_cv(perfil, resumen_profesional, bullets_por_trabajo, resumen_fit=resumen_fit)
    if not secciones:
        raise ValueError("No se puede generar el DOCX: no hay secciones con contenido.")

    doc = Document()

    # Configurar márgenes de página (0.75 pulgadas)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # 1. Encabezado del candidato
    nombre = perfil.get("nombre") or "Candidato/a"
    p_nombre = doc.add_paragraph()
    r_nombre = p_nombre.add_run(nombre)
    r_nombre.font.size = Pt(20)
    r_nombre.font.bold = True
    r_nombre.font.color.rgb = COLOR_TITULO
    p_nombre.paragraph_format.space_after = Pt(2)

    # Datos de contacto
    datos_contacto = [
        dato for dato in (perfil.get("ciudad"), perfil.get("email"), perfil.get("telefono"), perfil.get("linkedin")) if dato
    ]
    if datos_contacto:
        p_contacto = doc.add_paragraph()
        r_contacto = p_contacto.add_run("  ·  ".join(datos_contacto))
        r_contacto.font.size = Pt(9.5)
        r_contacto.font.color.rgb = COLOR_CONTACTO
        p_contacto.paragraph_format.space_after = Pt(8)

    # Subtítulo del Puesto
    p_puesto = doc.add_paragraph()
    r_puesto = p_puesto.add_run(f"CV Profesional — {puesto}")
    r_puesto.font.size = Pt(12)
    r_puesto.font.bold = True
    r_puesto.font.color.rgb = COLOR_SUBTITULO
    p_puesto.paragraph_format.space_after = Pt(10)

    # 2. Renderizado de secciones
    for seccion in secciones:
        # Título de Sección
        p_titulo = doc.add_paragraph()
        r_titulo = p_titulo.add_run(seccion["titulo"].upper())
        r_titulo.font.size = Pt(11)
        r_titulo.font.bold = True
        r_titulo.font.color.rgb = COLOR_SUBTITULO
        p_titulo.paragraph_format.space_before = Pt(8)
        p_titulo.paragraph_format.space_after = Pt(4)

        if seccion["tipo"] == "parrafo":
            for linea in seccion["contenido"].split("\n"):
                if linea.strip():
                    _agregar_parrafo_con_negritas(doc, linea)
        elif seccion["tipo"] == "trabajos":
            for trabajo in seccion["contenido"]:
                if trabajo["encabezado"]:
                    p_enc = doc.add_paragraph()
                    r_enc = p_enc.add_run(trabajo["encabezado"])
                    r_enc.font.bold = True
                    r_enc.font.size = Pt(10.5)
                    r_enc.font.color.rgb = COLOR_TEXTO
                    p_enc.paragraph_format.space_before = Pt(4)
                    p_enc.paragraph_format.space_after = Pt(2)
                for bullet in trabajo["bullets"]:
                    _agregar_parrafo_con_negritas(doc, bullet, es_bullet=True)
        elif seccion["tipo"] == "estudios":
            for estudio in seccion["contenido"]:
                if estudio["encabezado"]:
                    texto = estudio["encabezado"]
                    if estudio.get("tipo"):
                        texto = f"{texto} ({estudio['tipo']})"
                    p_est = doc.add_paragraph()
                    r_est = p_est.add_run(texto)
                    r_est.font.bold = True
                    r_est.font.size = Pt(10)
                    p_est.paragraph_format.space_after = Pt(2)
        elif seccion["tipo"] == "lista":
            for item in seccion["contenido"]:
                _agregar_parrafo_con_negritas(doc, item, es_bullet=True)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
